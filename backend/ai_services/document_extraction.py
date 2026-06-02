from dataclasses import dataclass
from pathlib import Path
import re

from .models import CandidateDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_UPLOAD_SIZE = 10 * 1024 * 1024
NO_EMBEDDED_TEXT_MESSAGE = (
    "PDF enthält keinen eingebetteten Text. Vermutlich Scan/Bild-PDF. "
    "OCR ist noch nicht aktiviert."
)
PDF_READ_ERROR_MESSAGE = "PDF konnte nicht gelesen werden."


@dataclass(frozen=True)
class ExtractionResult:
    status: str
    text: str = ""
    error: str = ""

    @property
    def text_length(self):
        return len(self.text)


def is_supported_document(filename):
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_candidate_document_text(document):
    suffix = Path(document.original_filename).suffix.lower()
    if suffix == ".txt":
        return _extract_txt(document.file.path)
    if suffix == ".pdf":
        return _extract_pdf(document.file.path)
    if suffix == ".docx":
        return _extract_docx(document.file.path)
    return ExtractionResult(
        status=CandidateDocument.ExtractionStatus.UNSUPPORTED,
        error="Dateityp wird nicht unterstützt.",
    )


def apply_extraction_result(document, result):
    document.extraction_status = result.status
    document.extracted_text = result.text
    document.extracted_text_length = result.text_length
    document.extraction_error = result.error
    document.save(
        update_fields=[
            "extraction_status",
            "extracted_text",
            "extracted_text_length",
            "extraction_error",
            "updated_at",
        ]
    )
    return document


def _extract_txt(path):
    try:
        raw = Path(path).read_bytes()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="ignore")
        return _text_result(text)
    except Exception:
        return ExtractionResult(
            status=CandidateDocument.ExtractionStatus.FAILED,
            error="TXT-Datei konnte nicht gelesen werden.",
        )


def _extract_pdf(path):
    try:
        from pypdf import PdfReader

        parts = []
        with open(path, "rb") as file:
            reader = PdfReader(file)
            for page in reader.pages:
                parts.append(page.extract_text() or "")
        return _text_result(
            "\n\n".join(parts),
            empty_error=NO_EMBEDDED_TEXT_MESSAGE,
        )
    except Exception:
        return ExtractionResult(
            status=CandidateDocument.ExtractionStatus.FAILED,
            error=PDF_READ_ERROR_MESSAGE,
        )


def _extract_docx(path):
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _text_result(
            "\n".join(parts),
            empty_error="DOCX-Datei enthält keinen auslesbaren Text.",
        )
    except Exception:
        return ExtractionResult(
            status=CandidateDocument.ExtractionStatus.FAILED,
            error="DOCX-Datei konnte nicht gelesen werden.",
        )


def _text_result(text, empty_error="Datei enthält keinen auslesbaren Text."):
    normalized_text = _normalize_text(text)
    if normalized_text:
        return ExtractionResult(
            status=CandidateDocument.ExtractionStatus.SUCCESS,
            text=normalized_text,
        )
    return ExtractionResult(
        status=CandidateDocument.ExtractionStatus.UNSUPPORTED,
        error=empty_error,
    )


def _normalize_text(text):
    lines = []
    for line in text.splitlines():
        normalized_line = re.sub(r"[ \t]+", " ", line).strip()
        if normalized_line:
            lines.append(normalized_line)
    return "\n".join(lines).strip()
